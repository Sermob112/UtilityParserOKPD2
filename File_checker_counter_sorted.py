import os
import re
from pathlib import Path
from typing import List

# ================== НАСТРОЙКИ ПО УМОЛЧАНИЮ ==================
DEFAULT_BASE = Path(r"C:\Users\Sergey\MAGA\Kimch\467 закупок")
DEFAULT_ORDER_TXT    = "purchase_numbers_from_excel.txt"
DEFAULT_OUTPUT_TXT   = "folder_structure_report.txt"
DEFAULT_OUTPUT_DOCX  = "folder_structure_report.docx"
DEFAULT_MISSING_TXT  = "missing_in_excel_order.txt"
# ============================================================

def read_numbers_in_order(order_file: Path) -> List[str]:
    nums = []
    with order_file.open("r", encoding="utf-8") as f:
        for line in f:
            s = line.strip()
            if s:
                nums.append(s)
    return nums

def digits_only(s: str) -> str:
    # правильный regex (без двойных слэшей)
    return re.sub(r"\D+", "", s, flags=re.UNICODE)

def build_no_patterns(num_key: str) -> List[re.Pattern]:
    # корректные паттерны с символом номера и аналогами
    markers = [r"№", r"Nº", r"N°", r"No", r"no", r"№\.?", r"N[oº°]\.?"]
    return [re.compile(rf"{m}\s*{re.escape(num_key)}", flags=re.IGNORECASE) for m in markers]

def name_matches_number(name: str, num_raw: str) -> bool:
    key = digits_only(num_raw)
    if not key:
        return False
    # 1) по цифрам-только
    if key in digits_only(name):
        return True
    # 2) явные «№ ...» и аналоги
    for p in build_no_patterns(key):
        if p.search(name):
            return True
    return False

def icon_for(name: str) -> str:
    ext = Path(name).suffix.lower()
    if ext in {".doc", ".docx"}: return "📄"
    if ext in {".pdf"}: return "📄"
    if ext in {".html", ".htm"}: return "🌐"
    if ext in {".xml"}: return "📄"
    if ext in {".xls", ".xlsx", ".csv"}: return "📊"
    if ext in {".ppt", ".pptx"}: return "📈"
    if ext in {".zip", ".rar", ".7z"}: return "📦"
    if ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp"}: return "🖼️"
    if ext in {".mp4", ".mkv", ".avi", ".mov", ".wmv"}: return "🎞️"
    if ext in {".mp3", ".wav", ".flac", ".aac", ".ogg"}: return "🎵"
    if ext in {".txt"}: return "🗒️"
    return "📄"

def make_docx_writer(docx_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn
    except Exception as e:
        print(f"[WARN] python-docx не найден: {e}")
        print("       Установите пакет: pip install python-docx")
        return None, None, None

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    try:
        font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except Exception:
        pass

    def add_header(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        except Exception:
            pass
        return p

    def add_item(text: str, level: int):
        p = doc.add_paragraph()
        try:
            p.paragraph_format.left_indent = Cm(0.6 * level)
        except Exception:
            pass
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        except Exception:
            pass
        return p

    return doc, add_header, add_item

def write_tree_txt(root_path: Path, out, indent: int = 0):
    try:
        items = sorted(os.listdir(root_path))
    except Exception as e:
        out.write(f"{'    '*indent}[ОШИБКА доступа: {e}]\n")
        return
    for item in items:
        p = root_path / item
        if p.is_dir():
            out.write("    " * indent + f"📁 {item}\n")
            write_tree_txt(p, out, indent + 1)
        else:
            out.write("    " * indent + f"{icon_for(item)} {item}\n")

def write_tree_docx(root_path: Path, add_item, indent: int = 0):
    try:
        items = sorted(os.listdir(root_path))
    except Exception as e:
        add_item(f"[ОШИБКА доступа: {e}]", indent)
        return
    for item in items:
        p = root_path / item
        if p.is_dir():
            add_item(f"📁 {item}", indent)
            write_tree_docx(p, add_item, indent + 1)
        else:
            add_item(f"{icon_for(item)} {item}", indent)

def main():
    script_dir = Path(__file__).resolve().parent
    order_path   = script_dir / DEFAULT_ORDER_TXT
    output_txt   = script_dir / DEFAULT_OUTPUT_TXT
    output_docx  = script_dir / DEFAULT_OUTPUT_DOCX
    missing_path = script_dir / DEFAULT_MISSING_TXT
    base_path    = DEFAULT_BASE

    if not order_path.exists():
        raise FileNotFoundError(f"Не найден '{order_path.name}'. Положите его рядом со скриптом: {script_dir}")
    if not base_path.exists():
        raise FileNotFoundError(f"Не найдена базовая директория:\n  {base_path}\nИзмените DEFAULT_BASE вверху файла.")

    nums_in_order = read_numbers_in_order(order_path)
    top_folders = [f for f in os.listdir(base_path) if (base_path / f).is_dir()]

    doc, add_header, add_item = make_docx_writer(output_docx)

    missing_numbers: List[str] = []
    blocks_written = 0

    with output_txt.open("w", encoding="utf-8") as out:
        for num in nums_in_order:
            matched = [folder for folder in top_folders if name_matches_number(folder, num)]
            if not matched:
                missing_numbers.append(num)
                continue

            out.write(f"\n Закупка № {num}\n")
            if doc is not None:
                add_header(f"Закупка № {num}")

            for folder in sorted(matched):
                folder_path = base_path / folder
                write_tree_txt(folder_path, out, indent=0)
                if doc is not None:
                    write_tree_docx(folder_path, add_item, indent=0)
                blocks_written += 1

    with missing_path.open("w", encoding="utf-8") as mf:
        for num in missing_numbers:
            mf.write(f"{num}\n")

    if doc is not None:
        try:
            doc.save(str(output_docx))
        except Exception as e:
            print(f"[WARN] Не удалось сохранить DOCX: {e}")

    print("[OK] Готово (исправленная версия).")
    print(f"  Номеров в TXT:       {len(nums_in_order)}")
    print(f"  Блоков записано:     {blocks_written}")
    print(f"  Не найдено номеров:  {len(missing_numbers)}")
    print(f"  TXT:                 {output_txt}")
    print(f"  DOCX:                {output_docx if doc is not None else '[python-docx не установлен]'}")
    print(f"  Не найдено (TXT):    {missing_path}")

if __name__ == "__main__":
    main()
