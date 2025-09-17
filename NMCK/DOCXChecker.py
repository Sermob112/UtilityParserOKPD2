# DOCChecker.py
import os
import re
import csv
import subprocess
from typing import Dict, List, Tuple, Set

# ================== НАСТРОЙКИ ==================
TARGET_DIR = r"C:\Users\Sergey\MAGA\Kimch\НМЦК_сборка\223- фз"
RECURSIVE = False         # обходить подпапки?
WRITE_SNIPPETS = True     # писать фрагменты в CSV
TARGET_ARTICLES = {"22", "34"}  # какие статьи считаем целевыми

# ================ ПАТТЕРНЫ ПОИСКА ==============
# «ст. 22», «ст22», «Ст.34», «ст 22», «статья 22», падежи «статьи/статье/статью/статьёй»
# + допускаем латинские look-alike буквы c/t в сокращении «ст.»
RE_ART_ANY = re.compile(
    r"""
    (?<!\w)
    (?:[сc][тt](?:\s*\.)? | стать(?:я|и|е|ю|ёй|ей))
    \s*
    ([0-9]{1,3}(?:\.\d+)*)    # номер статьи: 22, 34, 22.1, 34.2.5
    \b
    """,
    re.IGNORECASE | re.UNICODE | re.VERBOSE
)

# Приказы 567/639: не часть большего числа слева и
# НЕ должны быть дальше цифры (с учётом пробелов/.,,)
RE_ORDERS = re.compile(r"(?<!\d)(567|639)(?![\s.,]*\d)", re.UNICODE)

# Номер закупки из имени: приоритетно "Закупка № <19 цифр>", запасной — любой 19-значный
RE_NUM_FROM_TITLE = re.compile(r"Закупка\s*№\s*(\d{11})", re.UNICODE | re.IGNORECASE)
RE_ANY_19_DIGITS = re.compile(r"(?<!\d)(\d{11})(?!\d)", re.UNICODE)

# ================ УТИЛИТЫ ======================
def _normalize_ws(s: str) -> str:
    # NBSP / thin space / узкий NBSP / мягкий перенос -> обычный пробел/ничего
    return (s or "").replace("\u00A0", " ").replace("\u2009", " ").replace("\u202F", " ").replace("\u00AD", "")

def extract_purchase_number_from_filename(filename: str) -> str:
    m = RE_NUM_FROM_TITLE.search(filename)
    if m:
        return m.group(1)
    m2 = RE_ANY_19_DIGITS.search(filename)
    return m2.group(1) if m2 else "UNKNOWN"

def write_csv(path: str, rows: List[List[str]], header: List[str]):
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(header)
        w.writerows(rows)

def list_word_files(root: str, recursive: bool = False) -> List[str]:
    res = []
    if recursive:
        for dp, _, fns in os.walk(root):
            for fn in fns:
                low = fn.lower()
                if low.endswith(".docx") or (low.endswith(".doc") and not low.endswith(".docx")):
                    res.append(os.path.join(dp, fn))
    else:
        if not os.path.isdir(root):
            return []
        for fn in os.listdir(root):
            low = fn.lower()
            if low.endswith(".docx") or (low.endswith(".doc") and not low.endswith(".docx")):
                res.append(os.path.join(root, fn))
    return res

# ============ ЧТЕНИЕ DOCX / DOC ===============
def read_docx_texts(path: str) -> List[str]:
    """Возвращает список строк из абзацев и ячеек таблиц DOCX."""
    texts: List[str] = []
    try:
        import docx  # python-docx
    except Exception as e:
        print(f"[WARN] Не установлен python-docx для DOCX: {e}")
        return texts
    try:
        doc = docx.Document(path)
        # абзацы
        for p in doc.paragraphs:
            t = _normalize_ws(p.text)
            if t:
                texts.append(t)
        # таблицы
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = _normalize_ws(cell.text)
                    if t:
                        texts.append(t)
    except Exception as e:
        print(f"[WARN] Ошибка чтения DOCX: {path}\n  {e}")
    return texts

def read_doc_via_com(path: str) -> str:
    """Пробуем открыть .doc через MS Word (COM). Требуется установленный MS Word и pywin32."""
    try:
        import win32com.client  # pywin32
        from win32com.client import constants
    except Exception:
        return ""
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path)
        txt = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return txt
    except Exception as e:
        # Попробуем корректно закрыть Word, если открыт
        try:
            word.Quit()
        except Exception:
            pass
        print(f"[WARN] COM не смог прочитать DOC: {path}\n  {e}")
        return ""

def read_doc_via_textract(path: str) -> str:
    """Пробуем извлечь .doc через textract (нужны внешние зависимости: antiword/catdoc)."""
    try:
        import textract
    except Exception:
        return ""
    try:
        b = textract.process(path)
        return b.decode("utf-8", "ignore")
    except Exception as e:
        print(f"[WARN] textract не смог прочитать DOC: {path}\n  {e}")
        return ""

def read_doc_via_antiword(path: str) -> str:
    """Пробуем через утилиту antiword (если установлена и в PATH)."""
    try:
        # -m UTF-8.txt -> вывод в UTF-8, может отсутствовать; тогда просто без -m
        cmd = ["antiword", "-m", "UTF-8.txt", path]
        out = subprocess.check_output(cmd, stderr=subprocess.STDOUT)
        return out.decode("utf-8", "ignore")
    except Exception:
        try:
            out = subprocess.check_output(["antiword", path], stderr=subprocess.STDOUT)
            return out.decode("utf-8", "ignore")
        except Exception as e2:
            print(f"[WARN] antiword не смог прочитать DOC: {path}\n  {e2}")
            return ""

def read_doc_texts(path: str) -> List[str]:
    """Читает .doc/.docx и возвращает список текстовых блоков (для сниппетов)."""
    low = path.lower()
    if low.endswith(".docx"):
        return read_docx_texts(path)
    if low.endswith(".doc"):
        # .doc: пытаемся по очереди COM → textract → antiword
        txt = read_doc_via_com(path)
        if not txt:
            txt = read_doc_via_textract(path)
        if not txt:
            txt = read_doc_via_antiword(path)
        txt = _normalize_ws(txt)
        return [t for t in [txt] if t]
    return []

# ============== ПОИСК СОВПАДЕНИЙ ==============
def find_matches(texts: List[str]) -> Tuple[Set[str], Dict[str, List[str]], Dict[str, List[str]]]:
    """
    Возвращает:
      - found_articles: множество корней статей {'22', '34'} найденных в тексте
      - orders_snips: {'567': [сниппеты], '639': [сниппеты]}
      - art_snips: {'22': [сниппеты], '34': [сниппеты]}
    """
    found_articles: Set[str] = set()
    orders_snips: Dict[str, List[str]] = {"567": [], "639": []}
    art_snips: Dict[str, List[str]] = {}

    def snippet(s: str, m: re.Match, radius: int = 40) -> str:
        if not WRITE_SNIPPETS:
            return ""
        start = max(0, m.start() - radius)
        end = min(len(s), m.end() + radius)
        return s[start:end].replace("\n", " ")

    for raw in texts:
        s = _normalize_ws(str(raw))

        # Статьи (22/34 + подпункты)
        for m in RE_ART_ANY.finditer(s):
            art_num = m.group(1)
            root = art_num.split('.')[0]  # '22' из '22.1'
            if root in TARGET_ARTICLES:
                found_articles.add(root)
                art_snips.setdefault(root, []).append(snippet(s, m))

        # Приказы 567/639
        for m in RE_ORDERS.finditer(s):
            ord_num = m.group(1)
            orders_snips.setdefault(ord_num, []).append(snippet(s, m))

    # Уберём пустые ключи
    orders_snips = {k: v for k, v in orders_snips.items() if v}
    art_snips = {k: v for k, v in art_snips.items() if v}
    return found_articles, orders_snips, art_snips

# ===================== MAIN ====================
def main():
    files = list_word_files(TARGET_DIR, RECURSIVE)
    if not files:
        print(f"В папке нет DOC/DOCX-файлов: {TARGET_DIR}")
        return

    rows_found_any: List[List[str]] = []   # статья или приказ, или оба
    rows_none: List[List[str]] = []        # ни статьи, ни приказа
    rows_only_639: List[List[str]] = []    # отдельный отчёт: все, где есть 639

    total_files   = len(files)
    docx_total    = sum(1 for p in files if p.lower().endswith(".docx"))
    doc_total     = sum(1 for p in files if p.lower().endswith(".doc"))
    docx_read_ok  = 0
    doc_read_ok   = 0

    # статистика
    cnt_both = 0
    cnt_orders_only = 0
    cnt_articles_only = 0
    cnt_none = 0
    cnt_with_639 = 0

    for path in files:
        fname = os.path.basename(path)
        purchase = extract_purchase_number_from_filename(fname)

        texts = read_doc_texts(path)

        # учёт успешности чтения
        low = path.lower()
        if low.endswith(".docx") and texts:
            docx_read_ok += 1
        if low.endswith(".doc") and texts:
            doc_read_ok += 1

        found_articles, orders_snips, art_snips = find_matches(texts)

        # агрегаты для CSV
        articles_found = "|".join(sorted(found_articles)) if found_articles else ""
        orders_found   = "|".join(sorted(orders_snips.keys())) if orders_snips else ""

        # сниппеты (по 1–2 на тип)
        if WRITE_SNIPPETS:
            article_snippets = " || ".join(
                f"{k}: " + " | ".join(v[:2]) for k, v in sorted(art_snips.items())
            ) if art_snips else ""
            order_snippets = " || ".join(
                f"{k}: " + " | ".join(v[:2]) for k, v in sorted(orders_snips.items())
            ) if orders_snips else ""
        else:
            article_snippets = ""
            order_snippets = ""

        # статус для found_any / not_found_any
        if found_articles and orders_snips:
            status = "orders+article"
            cnt_both += 1
            rows_found_any.append([
                purchase, articles_found, orders_found, status,
                fname, path, article_snippets, order_snippets
            ])
        elif orders_snips:
            status = "orders"
            cnt_orders_only += 1
            rows_found_any.append([
                purchase, articles_found, orders_found, status,
                fname, path, article_snippets, order_snippets
            ])
        elif found_articles:
            status = "article"
            cnt_articles_only += 1
            rows_found_any.append([
                purchase, articles_found, orders_found, status,
                fname, path, article_snippets, order_snippets
            ])
        else:
            status = "missing"
            cnt_none += 1
            rows_none.append([purchase, fname, path])

        # --- ОТЧЁТ ТОЛЬКО ПО 639 ---
        if "639" in orders_snips:
            cnt_with_639 += 1
            order_639_snippets = " | ".join(orders_snips["639"][:2]) if WRITE_SNIPPETS else ""
            rows_only_639.append([
                purchase, articles_found, orders_found, status,
                fname, path, order_639_snippets
            ])

    # пишем отчёты
    base = os.path.dirname(os.path.abspath(__file__))
    write_csv(os.path.join(base, "found_any_doc.csv"),
              rows_found_any,
              header=["purchase_number", "articles_found", "orders_found", "status",
                      "file_name", "file_path", "article_snippets", "order_snippets"])
    write_csv(os.path.join(base, "not_found_any_doc.csv"),
              rows_none,
              header=["purchase_number", "file_name", "file_path"])
    write_csv(os.path.join(base, "only_639_doc.csv"),
              rows_only_639,
              header=["purchase_number", "articles_found", "orders_found", "status",
                      "file_name", "file_path", "order_639_snippets"])

    # итоги в консоль
    print(f"Всего DOC/DOCX: {total_files}")
    print(f"DOCX: прочитано {docx_read_ok} из {docx_total}")
    print(f"DOC : прочитано {doc_read_ok} из {doc_total}")
    print(f"Найдено (статья ИЛИ приказ ИЛИ оба): {len(rows_found_any)}")
    print(f"  — только приказ: {cnt_orders_only}")
    print(f"  — только статья: {cnt_articles_only}")
    print(f"  — статья+приказ: {cnt_both}")
    print(f"Ничего не найдено: {cnt_none}")
    print(f"Документов с приказом 639: {cnt_with_639}")
    print("Готово: found_any_doc.csv, not_found_any_doc.csv, only_639_doc.csv")


if __name__ == "__main__":
    main()
